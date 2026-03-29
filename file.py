import io
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor

def generate_word_quotation(q):
    doc = Document()
    # --- SET GLOBAL FONT TO CALIBRI ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # --- 1. A4 PAGE SETUP & NARROW MARGINS ---
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)
    section.header_distance = Cm(0.5)
    section.footer_distance = Cm(0.5)
    
    content_width = Cm(18.46)
    
    # --- 2. THE FIXED HEADER ---
    section.different_first_page_header_footer = True
    header = section.first_page_header
    htable = header.add_table(1, 2, width=content_width)
    
    htable.allow_autofit = False  
    htable.table_layout = 'fixed' 
    htable.columns[0].width = Cm(14.5) 
    htable.columns[1].width = Cm(3.96)
    htable.rows[0].cells[0].width = Cm(14.5)
    htable.rows[0].cells[1].width = Cm(3.96)

    for cell in htable.rows[0].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for node in ['top', 'left', 'bottom', 'right']:
            node_el = OxmlElement(f'w:{node}')
            node_el.set(qn('w:w'), '0')
            node_el.set(qn('w:type'), 'dxa')
            tcMar.append(node_el)
        tcPr.append(tcMar)
    
    left_cell = htable.rows[0].cells[0]
    left_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    header_text = left_cell.paragraphs[0]
    run_name = header_text.add_run("JIGSAW AFRICA WILDLIFE SAFARIS LTD (JAWS AFRICA)\n")
    run_name.font.bold = True
    run_name.font.size = Pt(17)
    run_reg = header_text.add_run("REGISTRATION PIN: P052221766X\n")
    run_reg.font.bold = True
    run_reg.font.size = Pt(11)
    
    details = (
    "Mercantile House, 2nd floor, Room 230, Koinange street, Nairobi, Kenya\n"
    "Emails: info@jawsafrica.com | Web: www.jawsafrica.com\n"
    "Mobile: +254 719899245 (KE) | +965 94067244 (KW) | +91 9496656977 (IN)"
)
    run_details = header_text.add_run(details)
    run_details.font.size = Pt(11)

    right_cell = htable.rows[0].cells[1]
    right_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    logo_para = right_cell.paragraphs[0]
    logo_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT 
    if os.path.exists("logo.png"):
        logo_para.add_run().add_picture("logo.png", width=Cm(3.5))

    # --- 3. DOUBLE SOLID LINE DIVIDER ---
    line_para = header.add_paragraph()
    p_obj = line_para._p
    pPr = p_obj.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'double') 
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)

    # --- 4. BODY CONTENT (UNDERLINED TITLE WITH GAP) ---
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(10) 
    
    run = title.add_run(f"Quotation for {q['client']}| {q['country']}")
    run.bold = True
    run.underline = True 
    run.font.size = Pt(14)
    run.font.name = 'Calibri'

    summary_lines = [
        f"TOUR CODE:   {q['code']}",
        f"{q['country'].upper()} SAFARI PRIVATE PACKAGE :  {q['pkg']}", # Dynamic country variable
        f"DATE:   FROM {q['start']} TO {q['end']}"
    ]

    for line in summary_lines:
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(2) 
        p.paragraph_format.line_spacing = 1.0

    # --- HELPER FOR STYLED HEADINGS ---
    def add_styled_heading(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Cm(0) 
        
        # Gold background shading
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'FFC000') 
        p._p.get_or_add_pPr().append(shading_elm)
        
        run = p.add_run(f" {text}") # Leading space for padding
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)

    def set_cell_grey(cell):
        """Sets a light grey background for table headers"""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'D9D9D9') 
        cell._tc.get_or_add_tcPr().append(shading_elm)

    # --- 5. ITINERARY TABLE ---
    add_styled_heading('ITINERARY PLAN')
    iti_table = doc.add_table(rows=1, cols=6)
    iti_table.style = 'Table Grid'
    iti_table.alignment = WD_ALIGN_PARAGRAPH.LEFT
    iti_table.width = content_width

    # SHIFT TABLE RIGHT: Accessing XML to set left indentation
    tbl_pr = iti_table._element.xpath('w:tblPr')[0]
    tbl_ind = OxmlElement('w:tblInd')
    tbl_ind.set(qn('w:w'), '100') 
    tbl_ind.set(qn('w:type'), 'dxa')
    tbl_pr.append(tbl_ind)

    hdrs = ["Day", "From", "To", "Activities", "Accommodation", "Meal Plan"]
    for i, h in enumerate(hdrs):
        cell = iti_table.rows[0].cells[i]
        cell.text = h
        set_cell_grey(cell) 
        p = cell.paragraphs[0]
        if p.runs:
            run = p.runs[0]
            run.bold = True
            run.font.name = 'Calibri'

    # Populating ITINERARY data with explicit mapping
    for row_data in q['iti']:
        row_cells = iti_table.add_row().cells
        row_cells[0].text = str(row_data.get("Day", ""))
        row_cells[1].text = str(row_data.get("From", ""))
        row_cells[2].text = str(row_data.get("To", ""))
        row_cells[3].text = str(row_data.get("Activities", ""))
        row_cells[4].text = str(row_data.get("Accommodation", ""))
        row_cells[5].text = str(row_data.get("Meal Plan", ""))
        
        # Ensure Calibri for all table content
        for cell in row_cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = 'Calibri'

    # --- NEW: DETAILED PRICE BREAKDOWN TABLE ---
    if q.get('price_table') is not None:
        add_styled_heading('DETAILED PRICE BREAKDOWN')
        p_table = doc.add_table(rows=1, cols=2)
        p_table.style = 'Table Grid'
        p_table.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_table.width = content_width

        p_tbl_pr = p_table._element.xpath('w:tblPr')[0]
        p_tbl_ind = OxmlElement('w:tblInd')
        p_tbl_ind.set(qn('w:w'), '100') 
        p_tbl_ind.set(qn('w:type'), 'dxa')
        p_tbl_pr.append(p_tbl_ind)

        h_cells = p_table.rows[0].cells
        h_cells[0].text = "Traveler Category"; h_cells[1].text = "Total Cost (USD)"
        for c in h_cells: 
            set_cell_grey(c)
            p = c.paragraphs[0]
            if p.runs:
                p.runs[0].bold = True
                p.runs[0].font.name = 'Calibri'

        for p_row in q['price_table']:
            cells = p_table.add_row().cells
            cells[0].text = str(p_row.get("Category", ""))
            cells[1].text = f"{p_row.get('Cost', 0):,.2f}"
            for c in cells:
                for p in c.paragraphs:
                    for r in p.runs: r.font.name = 'Calibri'

    # --- 6. TARIFF TABLE ---
    add_styled_heading('TARIFF IN USD')
    
    # PAX Logic updated for Total price only when children are present
    has_kids = q.get('children_count', 0) > 0
    t_table = doc.add_table(rows=2, cols=2 if has_kids else 3)
    t_table.style = 'Table Grid'
    t_table.alignment = WD_ALIGN_PARAGRAPH.LEFT
    t_table.width = content_width

    # SHIFT TABLE RIGHT
    t_tbl_pr = t_table._element.xpath('w:tblPr')[0]
    t_tbl_ind = OxmlElement('w:tblInd')
    t_tbl_ind.set(qn('w:w'), '100') 
    t_tbl_ind.set(qn('w:type'), 'dxa')
    t_tbl_pr.append(t_tbl_ind)

    # Header Row for Tariff
    t_hdrs = ["Item", "Total Cost"]
    if not has_kids: t_hdrs.append("Cost per Adult")
    
    for i, h in enumerate(t_hdrs):
        cell = t_table.rows[0].cells[i]
        cell.text = h
        set_cell_grey(cell) 
        p = cell.paragraphs[0]
        if p.runs:
            p.runs[0].bold = True
            p.runs[0].font.name = 'Calibri'

    tr = t_table.rows[1].cells
    total_participants = q['adults'] + q.get('children_count', 0)
    tr[0].text = f"{total_participants} PAX Safari"
    tr[1].text = f"{q['total']:,.0f}"
    
    # Handle cost per adult key only if it exists (Fix for KeyError: 'pp')
    if not has_kids and len(tr) > 2:
        tr[2].text = f"{q.get('pp', 0):,.0f}"
    
    for cell in tr:
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].font.name = 'Calibri'

    # --- TARIFF SUMMARY SENTENCE ---
    summary_text = (
        f"This price is for {q['adults']} adults and {q.get('children_count', 0)} children traveling in {q['vehicles']} private "
        f"safari vehicle(s) with accommodation in {q['accommodation_summary']} "
        f"as per the itinerary above."
    )
    
    if q.get('extras_summary'):
        summary_text += f" This price also includes {q['extras_summary']}."
    
    p_summary = doc.add_paragraph(summary_text)
    p_summary.paragraph_format.space_after = Pt(8)
    p_summary.paragraph_format.left_indent = Cm(0.12)
    
    for run in p_summary.runs:
        run.font.name = 'Calibri'
        run.italic = True 
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(0, 0, 0)
        
    # --- 7. DETAILED ITINERARY SECTION (JUSTIFIED) ---
    if q.get('detailed_iti'):
        add_styled_heading('DETAILED ITINERARY')
        
        for day_info in q['detailed_iti']:
            p_day = doc.add_paragraph()
            p_day.paragraph_format.space_before = Pt(10)
            p_day.paragraph_format.space_after = Pt(2)
            run_day = p_day.add_run(f"{day_info['day']} :-")
            run_day.bold = True
            run_day.underline = True
            run_day.font.name = 'Calibri'
            run_day.font.size = Pt(11)
            
            p_details = doc.add_paragraph(day_info['details'])
            p_details.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_details.paragraph_format.left_indent = Cm(0.12)
            p_details.paragraph_format.line_spacing = 1.15
            
            for run in p_details.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(10.5)

    # --- 8. UPDATED INCLUSIONS ---
    add_styled_heading('INCLUSIONS')
    inclusions = [
        "Mid-Range accommodation with meal plans as stated in the itinerary",
        "Transport in a private vehicle & use of customized 4WD Land cruisers with a Driver/Guide",
        "Game drive time of your choice, support for full day game drive at no extra cost",
        "Drinking water on safari and transfers",
        "All National Parks entrance fees and Taxes"
    ]
    for item in inclusions: doc.add_paragraph(item, style='List Bullet')
    
    # --- 9. UPDATED EXCLUSIONS ---
    add_styled_heading('EXCLUSIONS')
    exclusions = [
        "International/Domestic Flights/Visa and travel insurance",
        "Alcoholic drinks or soft drinks",
        "Tips 10 USD per day per person for the driver",
        "Items of personal nature"
    ]
    for item in exclusions: doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph("\n\nThank you for Choosing Jaws Africa").alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- 10. FOOTER CODE ---
    def setup_styled_footer(footer_obj):
        footer_obj.is_linked_to_previous = False
        f_para = footer_obj.paragraphs[0]
        f_para.clear()
        f_table = footer_obj.add_table(1, 2, width=content_width)
        f_table.allow_autofit = False
        f_table.columns[0].width = Cm(13.0); f_table.columns[1].width = Cm(5.46)
        
        l_run = f_table.rows[0].cells[0].paragraphs[0].add_run("www.jawsafrica.com")
        l_run.font.name = 'Calibri'; l_run.font.size = Pt(10)
        
        r_para = f_table.rows[0].cells[1].paragraphs[0]
        r_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT; r_para.add_run("Page ")
        
        f_p = r_para._p
        for tag in ["begin", "PAGE", "end"]:
            r = OxmlElement('w:r')
            if tag in ["begin", "end"]:
                fld = OxmlElement('w:fldChar'); fld.set(qn('w:fldCharType'), tag); r.append(fld)
            else:
                txt = OxmlElement('w:instrText'); txt.text = tag; r.append(txt)
            f_p.append(r)
        
        for run in r_para.runs: run.font.name = 'Calibri'; run.font.size = Pt(10)

    setup_styled_footer(section.first_page_footer)
    setup_styled_footer(section.footer)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()